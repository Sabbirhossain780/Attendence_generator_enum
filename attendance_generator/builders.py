import os
from io import BytesIO
import lxml.etree as etree
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DXA         = 635
PAGE_W_DXA  = 15840
PAGE_H_DXA  = 12240
MARGIN_DXA  = 720
CONTENT_DXA = PAGE_W_DXA - 2*MARGIN_DXA
COL_SN, COL_DATE, COL_PURP, COL_DATA, COL_WORK, COL_REM = 500, 1500, 6600, 600, 2400, 2800
assert COL_SN+COL_DATE+COL_PURP+COL_DATA+COL_WORK+COL_REM == CONTENT_DXA


def dxa_emu(d):
    return d * DXA


def _upsert_tblW(tbl_el, dxa):
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl_el.insert(0,tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.insert(0,tblW)
    tblW.set(qn("w:w"),str(dxa)); tblW.set(qn("w:type"),"dxa")


def set_tbl_grid(tbl_el, col_widths):
    for old in tbl_el.findall(qn("w:tblGrid")): tbl_el.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is not None: tblPr.addnext(grid)
    else: tbl_el.insert(0, grid)


def _upsert_tcW(tc_el, dxa):
    tcPr = tc_el.find(qn("w:tcPr"))
    if tcPr is None: tcPr = OxmlElement("w:tcPr"); tc_el.insert(0,tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.insert(0,tcW)
    tcW.set(qn("w:w"),str(dxa)); tcW.set(qn("w:type"),"dxa")


def set_cell_border(cell, bordered=True):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")): tcPr.remove(old)
    tcB = OxmlElement("w:tcBorders")
    for side in ("top","start","bottom","end"):
        el = OxmlElement(f"w:{side}")
        if bordered:
            el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"4")
            el.set(qn("w:space"),"0"); el.set(qn("w:color"),"000000")
        else:
            el.set(qn("w:val"),"none")
        tcB.append(el)
    tcPr.append(tcB)


def set_cell_valign(cell, val="center"):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")): tcPr.remove(old)
    v = OxmlElement("w:vAlign"); v.set(qn("w:val"),val); tcPr.append(v)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")): tcPr.remove(old)
    m = OxmlElement("w:tcMar")
    for s,v in (("top",top),("start",left),("bottom",bottom),("end",right)):
        el = OxmlElement(f"w:{s}"); el.set(qn("w:w"),str(v)); el.set(qn("w:type"),"dxa"); m.append(el)
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is not None: vAlign.addprevious(m)
    else: tcPr.append(m)


def set_no_table_borders(tbl_el):
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None: return
    for old in tblPr.findall(qn("w:tblBorders")): tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for s in ("top","start","bottom","end","insideH","insideV"):
        el = OxmlElement(f"w:{s}"); el.set(qn("w:val"),"none"); tblB.append(el)
    tblLook = tblPr.find(qn("w:tblLook"))
    if tblLook is not None: tblLook.addprevious(tblB)
    else: tblPr.append(tblB)


def set_row_height(row, dxa):
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")): trPr.remove(old)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),str(dxa)); trH.set(qn("w:hRule"),"atLeast"); trPr.append(trH)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        th = OxmlElement("w:tblHeader")
        th.set(qn("w:val"), "1")
        trPr.append(th)


def para_align(para, alignment="center"):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:jc")): pPr.remove(old)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"),alignment); pPr.append(jc)


def para_spacing(para, before=0, after=0, exact_line=True):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")): pPr.remove(old)
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    if exact_line:
        spc.set(qn("w:line"),     "240")
        spc.set(qn("w:lineRule"), "exact")
    jc = pPr.find(qn("w:jc"))
    if jc is not None: jc.addprevious(spc)
    else: pPr.append(spc)


def zero_para_spacing(para):
    para_spacing(para, before=0, after=0, exact_line=True)


def set_tab_stop(para, tab_type, pos_dxa):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:tabs")): pPr.remove(old)
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
    tab.set(qn("w:val"),tab_type); tab.set(qn("w:pos"),str(pos_dxa)); tabs.append(tab)
    pPr.insert(0,tabs)


def add_tab(para):
    r = OxmlElement("w:r"); t = OxmlElement("w:tab"); r.append(t); para._p.append(r)


def add_run(para, text, bold=False, size_pt=10, italic=False):
    run = para.add_run(text)
    run.bold=bold; run.italic=italic; run.font.size=Pt(size_pt); run.font.name="Arial"
    return run


def configure_cell(cell, width_dxa, text="", bold=False, size_pt=10,
                   center=False, top=80, bottom=80, left=120, right=120, bordered=True):
    _upsert_tcW(cell._tc, width_dxa)
    set_cell_border(cell, bordered)
    set_cell_margins(cell, top=top, bottom=bottom, left=left, right=right)
    set_cell_valign(cell)
    p = cell.paragraphs[0]
    zero_para_spacing(p)
    if center: para_align(p, "center")
    if text != "" or text == 0:
        add_run(p, str(text), bold=bold, size_pt=size_pt)


def purpose_cell(cell, width_dxa):
    _upsert_tcW(cell._tc, width_dxa)
    set_cell_border(cell); set_cell_margins(cell,top=80,bottom=80,left=120,right=60)
    set_cell_valign(cell)
    p = cell.paragraphs[0]; zero_para_spacing(p)
    B = "☐"
    txt = f"{B} Training   {B} Survey   {B} Base work   {B} Travel   {B} Office work"
    run = p.add_run(txt); run.font.size=Pt(9); run.font.name="Arial"


def data_cell(cell, has_data, width_dxa):
    _upsert_tcW(cell._tc, width_dxa)
    set_cell_border(cell); set_cell_margins(cell,top=80,bottom=80,left=40,right=40)
    set_cell_valign(cell)
    p = cell.paragraphs[0]; zero_para_spacing(p); para_align(p,"center")
    if has_data=="Yes":
        run=p.add_run("Yes"); run.bold=True; run.font.size=Pt(9); run.font.name="Arial"


def fix_zoom(doc):
    for zoom in doc.settings.element.findall(qn("w:zoom")):
        if not zoom.get(qn("w:percent")): zoom.set(qn("w:percent"),"100")


def register_logo(section, logo_path):
    if not logo_path or not os.path.exists(logo_path): return None
    with open(logo_path,"rb") as f: img=f.read()
    rId, _ = section.header.part.get_or_add_image(BytesIO(img))
    return rId


def logo_anchor_xml(rId):
    W,H=914400,484632
    return (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:drawing>'
        f'<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' behindDoc="0" distT="0" distB="0" distL="114935" distR="114935"'
        f' simplePos="0" locked="0" layoutInCell="1" allowOverlap="1" relativeHeight="2">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>-95250</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>-123825</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{W}" cy="{H}"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
        f'<wp:docPr id="1" name="Logo"/>'
        f'<wp:cNvGraphicFramePr>'
        f'<a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>'
        f'</wp:cNvGraphicFramePr>'
        f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr><pic:cNvPr id="1" name="Logo"/>'
        f'<pic:cNvPicPr><a:picLocks noChangeAspect="1"/></pic:cNvPicPr></pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="{rId}"/>'
        f'<a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{W}" cy="{H}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        f'</pic:pic></a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>'
    )


def build_header(section, logo_rId, enu_name, enu_id, period_start, period_end, org_line1, org_line2):
    section.different_first_page_header_footer = False
    hdr = section.header
    for p in list(hdr.paragraphs): p._element.getparent().remove(p._element)
    mid = CONTENT_DXA // 2

    def nhp():
        p_el = OxmlElement("w:p"); hdr._element.append(p_el)
        from docx.text.paragraph import Paragraph as P
        return P(p_el, hdr._element)

    p1 = nhp(); set_tab_stop(p1,"center",mid); zero_para_spacing(p1)
    if logo_rId: p1._p.append(etree.fromstring(logo_anchor_xml(logo_rId)))
    add_tab(p1); add_run(p1, org_line1, bold=True, size_pt=14)

    p2 = nhp(); set_tab_stop(p2,"center",mid); zero_para_spacing(p2)
    add_tab(p2); add_run(p2, org_line2, bold=True, size_pt=14)

    p3 = nhp(); para_align(p3,"center"); para_spacing(p3,before=0,after=60,exact_line=True)
    add_run(p3, "Attendance Sheet", size_pt=14)

    p4 = nhp(); zero_para_spacing(p4)
    add_run(p4, "Name: ",              bold=True, size_pt=11)
    add_run(p4, f"{enu_name} ({enu_id})",          size_pt=11)
    add_run(p4, "    Designation: ",   bold=True, size_pt=11)
    add_run(p4, "Enumerator",                      size_pt=11)
    add_run(p4, "    Mobile No.: ",    bold=True, size_pt=11)

    p5 = nhp(); para_spacing(p5, before=0, after=80, exact_line=True)
    add_run(p5, "Period  ", bold=True, size_pt=11)
    add_run(p5, f"{period_start}  to  {period_end}", size_pt=11)


def build_footer(section):
    ftr = section.footer
    for p in list(ftr.paragraphs): p._element.getparent().remove(p._element)

    def nfp():
        p_el = OxmlElement("w:p"); ftr._element.append(p_el)
        from docx.text.paragraph import Paragraph as P
        return P(p_el, ftr._element)

    auth = nfp(); para_spacing(auth, before=160, after=60, exact_line=True)
    add_run(auth, "Authorized by:", bold=True, size_pt=10)

    SIG  = ["Enumerator","FA","RA","DA"]
    n    = len(SIG); w = CONTENT_DXA // n
    wids = [w]*(n-1) + [CONTENT_DXA - w*(n-1)]

    from docx import Document
    tmp     = Document(); sig_tbl = tmp.add_table(rows=2, cols=n)
    _upsert_tblW(sig_tbl._tbl, CONTENT_DXA)
    set_tbl_grid(sig_tbl._tbl, wids)
    set_no_table_borders(sig_tbl._tbl)

    for i, (name, wdxa) in enumerate(zip(SIG, wids)):
        cell0 = sig_tbl.rows[0].cells[i]
        _upsert_tcW(cell0._tc, wdxa)
        set_cell_border(cell0, bordered=False)
        set_cell_margins(cell0, top=0, bottom=0, left=80, right=80)
        para_spacing(cell0.paragraphs[0], before=0, after=160, exact_line=True)
        for _ in range(2):
            np2 = cell0.add_paragraph(); para_spacing(np2, before=0, after=160, exact_line=True)
        sp = cell0.add_paragraph(); zero_para_spacing(sp)
        run = sp.add_run("________________________"); run.font.size=Pt(9); run.font.name="Arial"

        cell1 = sig_tbl.rows[1].cells[i]
        _upsert_tcW(cell1._tc, wdxa)
        set_cell_border(cell1, bordered=False)
        set_cell_margins(cell1, top=0, bottom=0, left=80, right=80)
        p1n = cell1.paragraphs[0]; zero_para_spacing(p1n)
        run1 = p1n.add_run(name)
        run1.bold=True; run1.font.size=Pt(9); run1.font.name="Arial"

    ftr._element.append(sig_tbl._tbl)


def build_attendance_table(doc, rows):
    HDR = [("S/N",COL_SN),("Date",COL_DATE),("Purpose",COL_PURP),
           ("Data",COL_DATA),("Working Place",COL_WORK),("Remarks",COL_REM)]
    tbl = doc.add_table(rows=len(rows)+1, cols=6)
    _upsert_tblW(tbl._tbl, CONTENT_DXA)
    set_tbl_grid(tbl._tbl, [COL_SN, COL_DATE, COL_PURP, COL_DATA, COL_WORK, COL_REM])

    hdr_row = tbl.rows[0]; set_row_height(hdr_row, 400)
    set_repeat_header(hdr_row)
    for ci,(txt,w) in enumerate(HDR):
        configure_cell(hdr_row.cells[ci], w, txt, bold=True, center=True, size_pt=10,
                       left=40 if ci==3 else 120, right=40 if ci==3 else 120)

    for ri, r in enumerate(rows):
        c = tbl.rows[ri+1].cells
        configure_cell(c[0], COL_SN,   r["sn"],           center=True, size_pt=9)
        configure_cell(c[1], COL_DATE, r["date"],          center=True, size_pt=9)
        purpose_cell(c[2], COL_PURP)
        data_cell(c[3], r["has_data"], COL_DATA)
        configure_cell(c[4], COL_WORK, r["working_place"], size_pt=9)
        configure_cell(c[5], COL_REM,  "",                 size_pt=9)


def build_summary_table(doc, enu_name, enu_id, survey_days_count=0):
    LCOL = int(CONTENT_DXA*0.38); MCOL = int(CONTENT_DXA*0.38); RCOL = CONTENT_DXA-LCOL-MCOL
    ROWS = [
        ("Name",              enu_name,                                      ""),
        ("ID No.",            enu_id,                                        ""),
        ("Period",            "",                                            ""),
        ("Survey Days",       str(survey_days_count) if survey_days_count else "", str(survey_days_count) if survey_days_count else ""),
        ("Training Days",     "",                                            ""),
        ("Travel Days",       "",                                            ""),
        ("Base Work Days",    "",                                            ""),
        ("Office Work Days",  "",                                            ""),
        ("Holiday/Off Days",  "",                                            ""),
        ("Weekend Days",      "",                                            ""),
        ("Total Working Days","",                                            ""),
        ("Remarks / Notes",   "",                                            ""),
    ]
    tbl = doc.add_table(rows=len(ROWS)+1, cols=3)
    _upsert_tblW(tbl._tbl, CONTENT_DXA)
    set_tbl_grid(tbl._tbl, [LCOL, MCOL, RCOL])

    brow = tbl.rows[0]; set_row_height(brow, 440)
    brow.cells[0].merge(brow.cells[1])
    configure_cell(brow.cells[0], LCOL+MCOL, "Enumerator Summary",
                   bold=True, center=True, size_pt=11, top=100, bottom=100)
    configure_cell(brow.cells[2], RCOL, "Days Count",
                   bold=True, center=True, size_pt=9.5, top=100, bottom=100)

    for ri,(label,value,dcount) in enumerate(ROWS):
        lc,vc,dc = tbl.rows[ri+1].cells
        configure_cell(lc, LCOL, label, bold=True, size_pt=9.5)
        configure_cell(vc, MCOL, str(value), size_pt=9.5)
        configure_cell(dc, RCOL, str(dcount), center=True, size_pt=9.5)
